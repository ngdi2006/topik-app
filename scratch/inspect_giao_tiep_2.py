import docx
import os
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

docx_path = r"E:\TOPIK-IBT\topik-app\DATA-EPS\ky-nang-giao-tiep (2).docx"
print("Exists:", os.path.exists(docx_path))

doc = docx.Document(docx_path)
print("Total paragraphs:", len(doc.paragraphs))
print("Total tables:", len(doc.tables))

with open(r"e:\TOPIK-IBT\topik-app\scratch\giao_tiep_2_dump.txt", "w", encoding="utf-8") as f:
    f.write("--- PARAGRAPHS ---\n")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            f.write(f"P {i}: {text}\n")

    if doc.tables:
        f.write("\n--- TABLES ---\n")
        table = doc.tables[0]
        for row_idx, row in enumerate(table.rows):
            row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            # De-duplicate cell text if cells are merged
            unique_cells = []
            for cell_text in row_text:
                if not unique_cells or unique_cells[-1] != cell_text:
                    unique_cells.append(cell_text)
            f.write(f"Row {row_idx}: {unique_cells}\n")

print("Dumped structure to scratch/giao_tiep_2_dump.txt")
