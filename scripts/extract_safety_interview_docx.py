"""Synchronize canonical P8 safety questions from SXCT-ATLD1.docx.

The DOCX is the editorial source of truth. Repeated rows for the same question
inside one topic are merged into one question with multiple suggested answers.
"""

from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = ROOT / "DATA-EPS" / "SXCT-ATLD1.docx"
OUTPUT_PATH = ROOT / "supabase" / "migrations" / "202608120001_sync_workplace_safety_p8_from_docx.sql"
TOPIC_RE = re.compile(r"^주제\s*(\d+)\.\s*(.+)$")
HANGUL_RE = re.compile(r"[가-힣]")
LEGACY_ID_SEEDS = {
    "위험작업은 왜 2인1조로 해야 할까요?": "위험작업은 왜 2 인1 조로 해야 할까요?",
    "2인1조 작업이 무엇인가요?": "2 인1 조 작업이 무엇인가요?",
}


def normalize(value: str) -> str:
    return " ".join(value.replace("\u00a0", " ").split()).strip()


def sql_literal(value: str) -> str:
    return "'" + value.replace("'", "''") + "'"


def group_for_topic(number: int) -> tuple[str, str]:
    if number <= 6:
        return "before_work", "Trước khi làm việc"
    if number <= 20:
        return "during_work", "Trong khi làm việc"
    if number <= 25:
        return "after_work", "Sau khi làm việc"
    return "incident_response", "Khi có sự cố"


def parse_topic_heading(value: str) -> tuple[int, str, str] | None:
    match = TOPIC_RE.match(value)
    if not match:
        return None
    number = int(match.group(1))
    title = match.group(2).strip()
    korean, separator, vietnamese = title.partition(":")
    return number, korean.strip(), vietnamese.strip().rstrip(".") if separator else ""


def parse_source() -> list[dict[str, object]]:
    document = Document(SOURCE_PATH)
    raw_rows: list[dict[str, object]] = []
    current_topic: tuple[int, str, str] | None = None

    for table in document.tables:
        for row in table.rows:
            cells = [
                [normalize(paragraph.text) for paragraph in cell.paragraphs if normalize(paragraph.text)]
                for cell in row.cells
            ]
            if not cells:
                continue

            heading = next((parse_topic_heading(value) for value in cells[0] if parse_topic_heading(value)), None)
            if heading:
                current_topic = heading
                continue
            if current_topic is None or len(cells) < 2 or not cells[0]:
                continue

            question_text = cells[0][0]
            vietnamese_meaning = cells[0][1] if len(cells[0]) > 1 else ""
            answers_ko: list[str] = []
            answers_vi: list[str] = []

            for value in cells[1]:
                if value.startswith("→") or not HANGUL_RE.search(value):
                    if answers_ko:
                        while len(answers_vi) < len(answers_ko):
                            answers_vi.append("")
                        answers_vi[-1] = value.lstrip("→").strip()
                else:
                    answers_ko.append(value)
                    answers_vi.append("")

            topic_number, topic_ko, topic_vi = current_topic
            group_id, group_label = group_for_topic(topic_number)
            raw_rows.append(
                {
                    "topic_number": topic_number,
                    "topic_ko": topic_ko,
                    "topic_vi": topic_vi,
                    "group_id": group_id,
                    "group_label": group_label,
                    "question_text": question_text,
                    "vietnamese_meaning": vietnamese_meaning,
                    "suggested_answers": answers_ko,
                    "suggested_answers_vi": answers_vi,
                }
            )

    merged: dict[tuple[int, str], dict[str, object]] = {}
    for row in raw_rows:
        key = (int(row["topic_number"]), str(row["question_text"]))
        target = merged.setdefault(key, {**row, "suggested_answers": [], "suggested_answers_vi": []})
        if not target["vietnamese_meaning"] and row["vietnamese_meaning"]:
            target["vietnamese_meaning"] = row["vietnamese_meaning"]
        target_answers = target["suggested_answers"]
        target_answers_vi = target["suggested_answers_vi"]
        assert isinstance(target_answers, list) and isinstance(target_answers_vi, list)
        for index, answer in enumerate(row["suggested_answers"]):
            if answer not in target_answers:
                target_answers.append(answer)
                translations = row["suggested_answers_vi"]
                target_answers_vi.append(translations[index] if index < len(translations) else "")

    rows = list(merged.values())
    rows.sort(key=lambda item: (int(item["topic_number"]), str(item["question_text"])))
    validate(rows)
    return rows


def validate(rows: list[dict[str, object]]) -> None:
    topics = {int(row["topic_number"]) for row in rows}
    if topics != set(range(1, 31)):
        raise RuntimeError(f"Expected topics 1..30, found {sorted(topics)}")
    if len(rows) != 111:
        raise RuntimeError(f"Expected 111 canonical questions, found {len(rows)}")
    for row in rows:
        answers = row["suggested_answers"]
        answers_vi = row["suggested_answers_vi"]
        if not row["vietnamese_meaning"]:
            raise RuntimeError(f"Missing Vietnamese question: {row['question_text']}")
        if not answers or len(answers) != len(answers_vi) or any(not value for value in answers_vi):
            raise RuntimeError(f"Incomplete answer pair: {row['question_text']}")


def build_sql(rows: list[dict[str, object]]) -> str:
    question_counts = Counter(str(row["question_text"]) for row in rows)
    occurrences: Counter[str] = Counter()
    values: list[str] = []

    for row in rows:
        question = str(row["question_text"])
        occurrences[question] += 1
        if question_counts[question] == 1 or occurrences[question] == 1:
            id_seed = f"safety:{LEGACY_ID_SEEDS.get(question, question)}"
        else:
            id_seed = f"safety:{row['topic_number']}:{question}"
        answers = json.dumps(row["suggested_answers"], ensure_ascii=False)
        answers_vi = json.dumps(row["suggested_answers_vi"], ensure_ascii=False)
        values.append(
            "(" + ", ".join(
                [
                    f"md5({sql_literal(id_seed)})::uuid",
                    sql_literal(question),
                    sql_literal(str(row["vietnamese_meaning"])),
                    f"{sql_literal(answers)}::jsonb",
                    f"{sql_literal(answers_vi)}::jsonb",
                    sql_literal(str(row["group_id"])),
                    str(row["topic_number"]),
                    sql_literal(str(row["topic_ko"])),
                    sql_literal(str(row["topic_vi"])),
                ]
            ) + ")"
        )

    return f"""-- Generated from DATA-EPS/SXCT-ATLD1.docx. Do not hand-edit canonical content.
ALTER TABLE public.interview_questions
    ADD COLUMN IF NOT EXISTS suggested_answers_vi jsonb;

CREATE TEMP TABLE canonical_workplace_safety_questions (
    id uuid PRIMARY KEY,
    question_text text NOT NULL,
    vietnamese_meaning text NOT NULL,
    suggested_answers jsonb NOT NULL,
    suggested_answers_vi jsonb NOT NULL,
    safety_group text NOT NULL,
    safety_topic_number integer NOT NULL,
    safety_topic_ko text NOT NULL,
    safety_topic_vi text NOT NULL
) ON COMMIT DROP;

INSERT INTO canonical_workplace_safety_questions VALUES
{',\n'.join(values)};

-- Update existing deterministic IDs in place so learner history remains attached.
UPDATE public.interview_questions AS target
SET category = 'An toàn lao động',
    question_text = source.question_text,
    vietnamese_meaning = source.vietnamese_meaning,
    suggested_answers = source.suggested_answers,
    suggested_answers_vi = source.suggested_answers_vi,
    industry = 'Sản xuất chế tạo',
    safety_group = source.safety_group,
    safety_topic_number = source.safety_topic_number,
    safety_topic_ko = source.safety_topic_ko,
    safety_topic_vi = source.safety_topic_vi,
    updated_at = now()
FROM canonical_workplace_safety_questions AS source
WHERE target.id = source.id;

INSERT INTO public.interview_questions (
    id, category, question_text, vietnamese_meaning, suggested_answers,
    suggested_answers_vi, industry, safety_group, safety_topic_number,
    safety_topic_ko, safety_topic_vi
)
SELECT id, 'An toàn lao động', question_text, vietnamese_meaning,
       suggested_answers, suggested_answers_vi, 'Sản xuất chế tạo', safety_group,
       safety_topic_number, safety_topic_ko, safety_topic_vi
FROM canonical_workplace_safety_questions
ON CONFLICT (id) DO UPDATE SET
    category = EXCLUDED.category,
    question_text = EXCLUDED.question_text,
    vietnamese_meaning = EXCLUDED.vietnamese_meaning,
    suggested_answers = EXCLUDED.suggested_answers,
    suggested_answers_vi = EXCLUDED.suggested_answers_vi,
    industry = EXCLUDED.industry,
    safety_group = EXCLUDED.safety_group,
    safety_topic_number = EXCLUDED.safety_topic_number,
    safety_topic_ko = EXCLUDED.safety_topic_ko,
    safety_topic_vi = EXCLUDED.safety_topic_vi,
    updated_at = now();
"""


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--check", action="store_true", help="Validate without writing the migration")
    args = parser.parse_args()
    rows = parse_source()
    if not args.check:
        OUTPUT_PATH.write_text(build_sql(rows), encoding="utf-8")
    counts = Counter(int(row["topic_number"]) for row in rows)
    print(f"Validated {len(rows)} questions across {len(counts)} topics")
    print("Per-topic:", ", ".join(f"{key}:{counts[key]}" for key in sorted(counts)))
    if not args.check:
        print(f"Created {OUTPUT_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
